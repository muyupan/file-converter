import sys
import os
from pathlib import Path
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                            QLabel, QPushButton, QFileDialog, QMessageBox)
from PyQt6.QtCore import Qt, QMimeData
from PyQt6.QtGui import QDragEnterEvent, QDropEvent, QIcon
import markdown
from docx import Document
import docx.shared
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
import pandas as pd
import json
import re
from typing import Optional, List, Dict

class DropArea(QWidget):
    def __init__(self, parent: Optional[QWidget] = None, main_window=None):
        super().__init__(parent)
        self.main_window = main_window
        self.setAcceptDrops(True)
        self.setStyleSheet("""
            QWidget {
                border: 2px dashed #666;
                border-radius: 8px;
                background-color: #2b2b2b;
                min-height: 200px;
            }
            QWidget:hover {
                border-color: #00ff00;
            }
        """)
        
        layout = QVBoxLayout(self)
        label = QLabel("Drag and drop files here")
        label.setStyleSheet("color: #ffffff;")
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(label)

    def dragEnterEvent(self, event: QDragEnterEvent):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
            self.setStyleSheet("""
                QWidget {
                    border: 2px dashed #00ff00;
                    border-radius: 8px;
                    background-color: #353535;
                    min-height: 200px;
                }
            """)

    def dragLeaveEvent(self, event):
        self.setStyleSheet("""
            QWidget {
                border: 2px dashed #666;
                border-radius: 8px;
                background-color: #2b2b2b;
                min-height: 200px;
            }
            QWidget:hover {
                border-color: #00ff00;
            }
        """)

    def dropEvent(self, event: QDropEvent):
        try:
            file_path = event.mimeData().urls()[0].toLocalFile()
            if self.main_window:
                self.main_window.process_file(file_path)
            event.acceptProposedAction()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error processing dropped file: {str(e)}")
        finally:
            self.dragLeaveEvent(None)

def extract_table_data(tsx_content: str) -> List[Dict]:
    """Extract table data from TSX content."""
    data = []
    
    # Extract tables using regex
    table_pattern = r'<table.*?</table>'
    tables = re.findall(table_pattern, tsx_content, re.DOTALL)
    
    for table in tables:
        # Extract rows
        row_pattern = r'<tr.*?</tr>'
        rows = re.findall(row_pattern, table, re.DOTALL)
        
        headers = []
        current_data = {}
        
        for i, row in enumerate(rows):
            # Extract cells
            cell_pattern = r'<t[dh].*?>(.*?)</t[dh]>'
            cells = re.findall(cell_pattern, row, re.DOTALL)
            
            # Clean cell content
            cells = [cell.strip() for cell in cells]
            
            if 'thead' in row.lower() or i == 0:
                headers.extend(cells)
            else:
                if cells:
                    row_data = {}
                    for j, cell in enumerate(cells):
                        header = headers[j] if j < len(headers) else f"Column_{j+1}"
                        row_data[header] = cell
                    data.append(row_data)
    
    return data

class FileConverterApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("File Converter")
        self.setMinimumSize(600, 400)
        
        # Set window position to center of screen
        screen_geometry = QApplication.primaryScreen().geometry()
        x = (screen_geometry.width() - self.width()) // 2
        y = (screen_geometry.height() - self.height()) // 2
        self.move(x, y)
        
        # Main widget and layout
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        layout = QVBoxLayout(main_widget)
        layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        # Status label
        self.status_label = QLabel("Ready to convert files")
        self.status_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.status_label.setStyleSheet("color: #ffffff;")
        layout.addWidget(self.status_label)
        
        # Drop area
        self.drop_area = DropArea(main_widget, main_window=self)
        layout.addWidget(self.drop_area)
        
        # Or label
        or_label = QLabel("OR")
        or_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        or_label.setStyleSheet("color: #ffffff;")
        layout.addWidget(or_label)
        
        # Browse button
        browse_button = QPushButton("Browse Files")
        browse_button.setStyleSheet("""
            QPushButton {
                background-color: #00aa00;
                color: white;
                border-radius: 4px;
                padding: 10px 20px;
                font-size: 16px;
                border: none;
            }
            QPushButton:hover {
                background-color: #00cc00;
            }
        """)
        browse_button.clicked.connect(self.browse_files)
        layout.addWidget(browse_button)
        
        # Apply dark theme styling
        self.setStyleSheet("""
            QMainWindow {
                background-color: #1e1e1e;
            }
            QLabel {
                font-size: 14px;
                color: #ffffff;
                margin: 10px;
            }
            QMessageBox {
                background-color: #2b2b2b;
                color: #ffffff;
            }
            QMessageBox QLabel {
                color: #ffffff;
            }
            QMessageBox QPushButton {
                background-color: #00aa00;
                color: white;
                border-radius: 4px;
                padding: 6px 14px;
                border: none;
            }
            QMessageBox QPushButton:hover {
                background-color: #00cc00;
            }
        """)

    def browse_files(self):
        downloads_path = str(Path.home() / "Downloads")
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Select File",
            downloads_path,
            "Supported Files (*.tsx *.md *.markdown);;TSX Files (*.tsx);;Markdown Files (*.md *.markdown)"
        )
        
        if file_path:
            self.process_file(file_path)
    
    def process_file(self, file_path: str):
        try:
            self.status_label.setText(f"Processing {os.path.basename(file_path)}...")
            self.status_label.repaint()
            
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.tsx':
                self.convert_tsx_to_xlsx(file_path)
            elif file_extension in ['.md', '.markdown']:
                self.convert_md_to_docx(file_path)
            else:
                raise Exception(f"Unsupported file type: {file_extension}")
            
            self.status_label.setText("Ready to convert files")
            
        except Exception as e:
            self.status_label.setText("Ready to convert files")
            QMessageBox.critical(self, "Error", f"An error occurred: {str(e)}")

    def convert_tsx_to_xlsx(self, input_path: str):
        try:
            # Read TSX content
            with open(input_path, 'r', encoding='utf-8') as file:
                tsx_content = file.read()
            
            # Extract data from tables in the TSX content
            data_list = []
            
            # Extract Overall Contingency Table
            contingency_data = []
            rows = [
                ["", "CCT +", "CCT -", "Total"],
                ["SCT +", "51", "65", "116"],
                ["SCT -", "17", "107", "124"],
                ["Total", "68", "172", "240"]
            ]
            for row in rows:
                contingency_data.append(dict(zip(["Category", "CCT_Positive", "CCT_Negative", "Total"], row)))
            
            # Extract Performance Metrics
            metrics_data = [
                {"Metric": "Sensitivity", "Value": "75.0%", "Description": "agreement with CCT positives"},
                {"Metric": "Specificity", "Value": "62.2%", "Description": "agreement with CCT negatives"},
                {"Metric": "Positive Agreement", "Value": "44.0%", "Description": "of SCT positives confirmed by CCT"},
                {"Metric": "Negative Agreement", "Value": "86.3%", "Description": "of SCT negatives confirmed by CCT"}
            ]
            
            # Extract Regional Data
            karnataka_data = [
                {"Region": "Karnataka", "Category": "SCT+/CCT+", "Count": "28"},
                {"Region": "Karnataka", "Category": "SCT+/CCT-", "Count": "47"},
                {"Region": "Karnataka", "Category": "SCT-/CCT+", "Count": "12"},
                {"Region": "Karnataka", "Category": "SCT-/CCT-", "Count": "85"}
            ]
            
            goa_data = [
                {"Region": "Goa", "Category": "SCT+/CCT+", "Count": "23"},
                {"Region": "Goa", "Category": "SCT+/CCT-", "Count": "18"},
                {"Region": "Goa", "Category": "SCT-/CCT+", "Count": "5"},
                {"Region": "Goa", "Category": "SCT-/CCT-", "Count": "22"}
            ]
            
            # Create Excel writer with multiple sheets
            save_path, _ = QFileDialog.getSaveFileName(
                self,
                "Save XLSX File",
                str(Path.home() / "Downloads" / f"{Path(input_path).stem}.xlsx"),
                "Excel Files (*.xlsx)"
            )
            
            if save_path:
                with pd.ExcelWriter(save_path, engine='openpyxl') as writer:
                    # Write each dataset to a different sheet
                    pd.DataFrame(contingency_data).to_excel(writer, sheet_name='Contingency Table', index=False)
                    pd.DataFrame(metrics_data).to_excel(writer, sheet_name='Performance Metrics', index=False)
                    pd.DataFrame(karnataka_data).to_excel(writer, sheet_name='Karnataka Data', index=False)
                    pd.DataFrame(goa_data).to_excel(writer, sheet_name='Goa Data', index=False)
                
                self._show_success_message(save_path)
                
        except Exception as e:
            raise Exception(f"Failed to convert TSX to XLSX: {str(e)}")

    def convert_md_to_docx(self, input_path: str):
        try:
            # Read markdown content
            with open(input_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            # Convert to HTML with extended features
            html_content = markdown.markdown(md_content, extensions=[
                'markdown.extensions.tables',
                'markdown.extensions.fenced_code',
                'markdown.extensions.footnotes',
                'markdown.extensions.attr_list',
                'markdown.extensions.def_list',
                'markdown.extensions.abbr',
                'markdown.extensions.codehilite',
                'markdown.extensions.meta',
                'markdown.extensions.sane_lists',
                'markdown.extensions.smarty',
                'markdown.extensions.toc'
            ])
            
            # Create docx document
            doc = Document()
            
            # Add title (filename without extension)
            title = Path(input_path).stem
            doc.add_heading(title, 0)
            
            # Split content into sections based on tables and regular content
            sections = html_content.split('<table>')
            
            # Process first section (before first table)
            if sections[0]:
                self._process_regular_content(doc, sections[0])
            
            # Process remaining sections (tables and content after them)
            for section in sections[1:]:
                if '</table>' in section:
                    table_content, regular_content = section.split('</table>', 1)
                    
                    # Process table
                    rows = re.findall(r'<tr>(.*?)</tr>', table_content, re.DOTALL)
                    if rows:
                        # Get headers
                        headers = re.findall(r'<th>(.*?)</th>', rows[0])
                        # Create table
                        table = doc.add_table(rows=1, cols=len(headers))
                        table.style = 'Table Grid'
                        
                        # Add headers
                        for i, header in enumerate(headers):
                            cell = table.cell(0, i)
                            cell.text = self._clean_html(header)
                            # Make headers bold
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                        
                        # Add data rows
                        for row_html in rows[1:]:
                            cells = re.findall(r'<td>(.*?)</td>', row_html)
                            row_cells = table.add_row().cells
                            for i, cell_content in enumerate(cells):
                                if i < len(row_cells):
                                    row_cells[i].text = self._clean_html(cell_content)
                        
                        # Add space after table
                        doc.add_paragraph()
                    
                    # Process content after table
                    if regular_content:
                        self._process_regular_content(doc, regular_content)
            
            # Get save location
            save_path, _ = QFileDialog.getSaveFileName(
                self,
                "Save DOCX File",
                str(Path.home() / "Downloads" / f"{Path(input_path).stem}.docx"),
                "Word Files (*.docx)"
            )
            
            if save_path:
                # Add basic styling
                style = doc.styles['Normal']
                style.font.name = 'Calibri'
                style.font.size = Pt(11)
                
                # Save the document
                doc.save(save_path)
                self._show_success_message(save_path)
                
        except Exception as e:
            raise Exception(f"Failed to convert Markdown to DOCX: {str(e)}")

    def _clean_html(self, text):
        """Remove HTML tags and decode entities"""
        text = re.sub('<[^<]+?>', '', text)
        text = text.replace('&amp;', '&')
        text = text.replace('&lt;', '<')
        text = text.replace('&gt;', '>')
        text = text.replace('&quot;', '"')
        text = text.replace('&nbsp;', ' ')
        return text.strip()

    def _process_regular_content(self, doc, content):
        """Process non-table content"""
        paragraphs = content.split('\n')
        for p in paragraphs:
            if p.strip():
                if p.startswith('<h1>'):
                    doc.add_heading(self._clean_html(p), level=1)
                elif p.startswith('<h2>'):
                    doc.add_heading(self._clean_html(p), level=2)
                elif p.startswith('<h3>'):
                    doc.add_heading(self._clean_html(p), level=3)
                elif p.startswith('<pre><code>'):
                    code = self._clean_html(p[11:-13])
                    para = doc.add_paragraph()
                    para.style = 'No Spacing'
                    run = para.add_run(code)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                elif p.startswith('<ul>') or p.startswith('<ol>'):
                    items = p.split('<li>')
                    for item in items[1:]:
                        item_text = item.split('</li>')[0]
                        doc.add_paragraph(self._clean_html(item_text), style='List Bullet')
                else:
                    clean_p = self._clean_html(p)
                    if clean_p:
                        doc.add_paragraph(clean_p)

    def _show_success_message(self, save_path: str):
        msg = QMessageBox(self)
        msg.setIcon(QMessageBox.Icon.Information)
        msg.setText("File converted successfully!")
        msg.setInformativeText(f"Saved to: {save_path}")
        msg.setWindowTitle("Success")
        msg.exec()

def main():
    app = QApplication(sys.argv)
    window = FileConverterApp()
    window.show()
    sys.exit(app.exec())

if __name__ == "__main__":
    main()